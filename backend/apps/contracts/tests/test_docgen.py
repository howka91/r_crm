"""Tests for Phase 5.3 — docgen pipeline and generate-pdf endpoint.

Covers:
    * services.snapshot.build_context — keys and fallbacks
    * services.qr — PNG + data-URI shape
    * services.docgen — placeholder substitution, unknown-key handling,
      actual WeasyPrint PDF attachment.
    * POST /contracts/:id/generate-pdf/ — happy path, template-missing,
      unknown-placeholder, blocked-on-signed.
    * ContractTemplate CRUD — global vs per-project gate, placeholders
      validation.
"""
from __future__ import annotations

import io
from decimal import Decimal

import pytest
from django.urls import reverse
from rest_framework import status

from apps.contracts.models import Contract, ContractTemplate
from apps.contracts.services import docgen as docgen_svc
from apps.contracts.services.qr import make_qr_data_uri, make_qr_png
from apps.contracts.services.snapshot import build_context
from apps.contracts.tests.factories import ContractFactory, ContractTemplateFactory
from apps.core.permission_tree import default_permissions
from apps.users.tests.factories import RoleFactory, StaffFactory


def _scoped(*keys: str) -> dict[str, bool]:
    perms = default_permissions(False)
    for k in keys:
        parts = k.split(".")
        for i in range(1, len(parts) + 1):
            perms[".".join(parts[:i])] = True
    return perms


def _superuser(api_client):
    admin = StaffFactory(is_superuser=True, is_staff=True, password="x12345678")
    api_client.force_authenticate(admin)
    return admin


# --- snapshot -----------------------------------------------------------


@pytest.mark.django_db
class TestBuildContext:
    def test_root_keys_present(self):
        contract = ContractFactory()
        ctx = build_context(contract)
        # All documented roots appear.
        assert set(ctx) >= {
            "contract", "apartment", "project", "developer",
            "calculation", "signer", "client", "now", "today",
        }

    def test_client_reached_via_signer(self):
        contract = ContractFactory()
        ctx = build_context(contract)
        assert ctx["signer"] is contract.signer
        assert ctx["client"] is contract.signer.client

    def test_missing_calculation_yields_none(self):
        contract = ContractFactory()
        contract.calculation = None
        contract.save(update_fields=["calculation"])
        ctx = build_context(contract)
        assert ctx["calculation"] is None


# --- qr ------------------------------------------------------------------


class TestQR:
    def test_png_starts_with_signature(self):
        png = make_qr_png("hello")
        assert png[:4] == b"\x89PNG"

    def test_data_uri_shape(self):
        uri = make_qr_data_uri("hello")
        assert uri.startswith("data:image/png;base64,")
        assert len(uri) > 100


# --- docgen service -----------------------------------------------------


@pytest.mark.django_db
class TestGeneratePdf:
    def _tpl(self, **kw) -> ContractTemplate:
        return ContractTemplateFactory(
            body=kw.pop(
                "body",
                "<html><body><p>Договор №{{contract_number}} от {{date}}.</p>"
                "<p>Клиент: {{client_name}}</p><p>Квартира: {{apt_number}}</p>"
                "</body></html>",
            ),
            placeholders=kw.pop(
                "placeholders",
                [
                    {"key": "contract_number", "path": "contract.contract_number", "label": "№"},
                    {"key": "date", "path": "contract.date", "label": "Дата"},
                    {"key": "client_name", "path": "client.full_name", "label": "ФИО"},
                    {"key": "apt_number", "path": "apartment.number", "label": "Кв."},
                ],
            ),
            **kw,
        )

    def test_happy_path_attaches_pdf(self):
        contract = ContractFactory(contract_number="ЯМ-00001")
        contract.template = self._tpl()
        contract.save(update_fields=["template"])

        result = docgen_svc.generate_pdf(contract)
        contract.refresh_from_db()

        # File attached, PDF magic bytes present.
        assert contract.file, "Expected Contract.file to be populated"
        with contract.file.open("rb") as f:
            assert f.read(4) == b"%PDF"
        assert result.pdf_size > 0
        # document snapshot persisted.
        assert contract.document["template_title"] == contract.template.title
        assert contract.document["values"]["contract_number"] == "ЯМ-00001"

    def test_template_not_set_raises(self):
        contract = ContractFactory(template=None)
        with pytest.raises(docgen_svc.TemplateNotSet):
            docgen_svc.generate_pdf(contract)

    def test_unknown_placeholder_strict_raises(self):
        contract = ContractFactory()
        contract.template = self._tpl(
            body="<p>{{contract_number}} / {{mystery_key}}</p>",
            placeholders=[
                {"key": "contract_number", "path": "contract.contract_number", "label": "№"},
            ],
        )
        contract.save(update_fields=["template"])
        with pytest.raises(docgen_svc.UnknownPlaceholder) as exc:
            docgen_svc.generate_pdf(contract)
        assert "mystery_key" in exc.value.unknown

    def test_unknown_placeholder_lenient_passes_through(self):
        contract = ContractFactory()
        contract.template = self._tpl(
            body="<p>{{contract_number}} / {{mystery_key}}</p>",
            placeholders=[
                {"key": "contract_number", "path": "contract.contract_number", "label": "№"},
            ],
        )
        contract.save(update_fields=["template"])
        result = docgen_svc.generate_pdf(contract, strict=False)
        assert result.pdf_size > 0

    def test_missing_path_resolves_to_empty(self):
        """Paths that don't exist on the tree should render as "", not crash."""
        contract = ContractFactory()
        contract.template = self._tpl(
            body="<p>[{{bogus}}]</p>",
            placeholders=[
                {"key": "bogus", "path": "contract.nonexistent.field", "label": "x"},
            ],
        )
        contract.save(update_fields=["template"])
        result = docgen_svc.generate_pdf(contract)
        assert result.filled["bogus"] == ""


# --- /generate-pdf/ endpoint --------------------------------------------


@pytest.mark.django_db
class TestGeneratePdfEndpoint:
    def test_happy_path(self, api_client):
        _superuser(api_client)
        contract = ContractFactory(contract_number="ЯМ-00010")
        contract.template = ContractTemplateFactory(
            body="<p>{{n}}</p>",
            placeholders=[
                {"key": "n", "path": "contract.contract_number", "label": "№"},
            ],
        )
        contract.save(update_fields=["template"])

        resp = api_client.post(
            reverse("contract-generate-pdf", args=[contract.id]),
        )
        assert resp.status_code == status.HTTP_200_OK, resp.data
        assert resp.data["pdf_size"] > 0
        assert resp.data["filled"]["n"] == "ЯМ-00010"

    def test_template_not_set_returns_400(self, api_client):
        _superuser(api_client)
        contract = ContractFactory(template=None)
        resp = api_client.post(
            reverse("contract-generate-pdf", args=[contract.id]),
        )
        assert resp.status_code == status.HTTP_400_BAD_REQUEST
        assert resp.data["code"] == "template_not_set"

    def test_unknown_placeholder_returns_400(self, api_client):
        _superuser(api_client)
        contract = ContractFactory()
        contract.template = ContractTemplateFactory(
            body="<p>{{missing}}</p>",
            placeholders=[],
        )
        contract.save(update_fields=["template"])
        resp = api_client.post(
            reverse("contract-generate-pdf", args=[contract.id]),
        )
        assert resp.status_code == status.HTTP_400_BAD_REQUEST
        assert resp.data["code"] == "unknown_placeholder"
        assert "missing" in resp.data["unknown"]

    def test_blocked_on_signed_contract(self, api_client):
        _superuser(api_client)
        contract = ContractFactory(is_signed=True)
        contract.template = ContractTemplateFactory()
        contract.save(update_fields=["template"])
        resp = api_client.post(
            reverse("contract-generate-pdf", args=[contract.id]),
        )
        assert resp.status_code == status.HTTP_409_CONFLICT


# --- ContractTemplate CRUD with global-template gate --------------------


@pytest.mark.django_db
class TestContractTemplateGate:
    url_list = reverse("contract-template-list")

    def test_scoped_user_can_create_project_template(self, api_client):
        from apps.objects.tests.factories import ProjectFactory
        role = RoleFactory(
            code="tpl-creator-project",
            permissions=_scoped(
                "references.templates.view",
                "references.templates.create",
            ),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        project = ProjectFactory()
        resp = api_client.post(
            self.url_list,
            {
                "title": "Scoped template",
                "body": "<p>x</p>",
                "project": project.id,
                "placeholders": [],
            },
            format="json",
        )
        assert resp.status_code == status.HTTP_201_CREATED, resp.data

    def test_scoped_user_blocked_from_global_template(self, api_client):
        role = RoleFactory(
            code="tpl-creator-noglobal",
            permissions=_scoped(
                "references.templates.view",
                "references.templates.create",
            ),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        resp = api_client.post(
            self.url_list,
            {"title": "Global attempt", "body": "<p>x</p>", "placeholders": []},
            format="json",
        )
        assert resp.status_code == status.HTTP_403_FORBIDDEN
        assert resp.data["code"] == "global_template_forbidden"

    def test_manage_global_role_can_create_global(self, api_client):
        role = RoleFactory(
            code="tpl-creator-global",
            permissions=_scoped(
                "references.templates.view",
                "references.templates.create",
                "references.templates.manage_global",
            ),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        resp = api_client.post(
            self.url_list,
            {"title": "Global", "body": "<p>x</p>", "placeholders": []},
            format="json",
        )
        assert resp.status_code == status.HTTP_201_CREATED, resp.data
        assert resp.data["is_global"] is True

    def test_superuser_bypasses_global_gate(self, api_client):
        _superuser(api_client)
        resp = api_client.post(
            self.url_list,
            {"title": "Admin global", "body": "<p>x</p>", "placeholders": []},
            format="json",
        )
        assert resp.status_code == status.HTTP_201_CREATED, resp.data

    def test_placeholder_validation_rejects_missing_path(self, api_client):
        _superuser(api_client)
        resp = api_client.post(
            self.url_list,
            {
                "title": "Broken",
                "body": "<p>{{x}}</p>",
                "placeholders": [{"key": "x", "label": "no path"}],
            },
            format="json",
        )
        assert resp.status_code == status.HTTP_400_BAD_REQUEST

    def test_placeholder_validation_rejects_duplicate_keys(self, api_client):
        _superuser(api_client)
        resp = api_client.post(
            self.url_list,
            {
                "title": "Dup",
                "body": "<p>{{x}}</p>",
                "placeholders": [
                    {"key": "x", "path": "contract.id", "label": "a"},
                    {"key": "x", "path": "contract.date", "label": "b"},
                ],
            },
            format="json",
        )
        assert resp.status_code == status.HTTP_400_BAD_REQUEST


# --- _BASE_CSS prelude --------------------------------------------------


@pytest.mark.django_db
class TestBaseCss:
    """`_BASE_CSS` is prepended to every template body before it hits
    WeasyPrint. Keeping its rules stable is how the editor preview and
    the rendered PDF agree on page-breaks + image alignment — if someone
    edits the constant, these assertions should fail loudly."""

    def test_page_rule_sets_A4_and_margins(self):
        assert "@page" in docgen_svc._BASE_CSS
        assert "size: A4" in docgen_svc._BASE_CSS
        assert "18mm 15mm" in docgen_svc._BASE_CSS

    def test_page_break_class_emits_break_before(self):
        """`.page-break` is how the Tiptap PageBreak node maps into print
        CSS. It must keep both the modern `break-before` and the legacy
        `page-break-before` for WeasyPrint compatibility."""
        assert ".page-break" in docgen_svc._BASE_CSS
        assert "break-before: page" in docgen_svc._BASE_CSS
        assert "page-break-before: always" in docgen_svc._BASE_CSS
        # Clear float context before the break — logos floated left
        # must not leak into the next page.
        assert "clear: both" in docgen_svc._BASE_CSS

    def test_img_data_align_rules(self):
        """`data-align` on images is the contract between the Tiptap
        AlignedImage extension and the PDF. All four layouts must be
        represented so the PDF mirrors the editor."""
        css = docgen_svc._BASE_CSS
        assert 'img[data-align="left"]' in css
        assert 'img[data-align="right"]' in css
        assert 'img[data-align="center"]' in css
        assert "float: left" in css
        assert "float: right" in css
        assert "margin: 4px auto" in css  # centered image

    def test_filled_body_reaches_renderer(self, monkeypatch):
        """Integration check — the HTML handed to `_render_pdf` is the
        template body with placeholders substituted, *before* `_BASE_CSS`
        is prepended (the prepend happens inside `_render_pdf` itself,
        so this captured value sees only the body). Verifies the
        substitution pipeline + that the page-break div survives."""
        captured: dict[str, str] = {}

        def fake_render(html: str) -> bytes:
            captured["html"] = html
            return b"%PDF-FAKE"

        monkeypatch.setattr(docgen_svc, "_render_pdf", fake_render)

        contract = ContractFactory(contract_number="CSS-TEST-1")
        contract.template = ContractTemplateFactory(
            body='<div class="page-break"></div><p>{{n}}</p>',
            placeholders=[
                {"key": "n", "path": "contract.contract_number", "label": "№"},
            ],
        )
        contract.save(update_fields=["template"])
        docgen_svc.generate_pdf(contract)

        html = captured["html"]
        # Placeholder substitution happened.
        assert "CSS-TEST-1" in html
        assert "{{n}}" not in html
        # The page-break div survives unchanged — `_render_pdf` wires
        # the actual break via `_BASE_CSS` on the way into WeasyPrint.
        assert '<div class="page-break"></div>' in html


# --- ContractTemplate upload-image action -------------------------------


# 1x1 transparent PNG — smallest possible valid PNG, handy for fixtures.
_PNG_1x1 = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\rIDATx\x9cc\x00"
    b"\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82"
)


@pytest.mark.django_db
class TestContractTemplateUploadImage:
    """`POST /contract-templates/upload-image/` — the endpoint the Tiptap
    editor hits when the admin inserts a logo. Image bytes land on
    default_storage; the response URL is what ends up in `<img src>`."""

    url = reverse("contract-template-upload-image")

    def _upload(self, api_client, **kw):
        from django.core.files.uploadedfile import SimpleUploadedFile
        content = kw.pop("content", _PNG_1x1)
        content_type = kw.pop("content_type", "image/png")
        name = kw.pop("name", "logo.png")
        file = SimpleUploadedFile(name, content, content_type=content_type)
        return api_client.post(self.url, {"file": file}, format="multipart")

    def test_happy_path_returns_url_and_metadata(self, api_client):
        _superuser(api_client)
        resp = self._upload(api_client)
        assert resp.status_code == status.HTTP_201_CREATED, resp.data
        assert resp.data["url"].startswith("/media/")
        assert "/contract_templates/images/" in resp.data["url"]
        assert resp.data["url"].endswith(".png")
        assert resp.data["content_type"] == "image/png"
        assert resp.data["size"] == len(_PNG_1x1)
        # Filename is a uuid — no leak of the client-supplied name.
        assert resp.data["filename"] != "logo.png"

    def test_unauthenticated_rejected(self, api_client):
        resp = self._upload(api_client)
        assert resp.status_code in (
            status.HTTP_401_UNAUTHORIZED,
            status.HTTP_403_FORBIDDEN,
        )

    def test_viewer_without_create_or_edit_blocked(self, api_client):
        """Plain `references.templates.view` shouldn't grant uploads —
        the gate specifically asks for create OR edit."""
        role = RoleFactory(
            code="tpl-view-only",
            permissions=_scoped("references.templates.view"),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        resp = self._upload(api_client)
        assert resp.status_code == status.HTTP_403_FORBIDDEN

    def test_editor_can_upload(self, api_client):
        role = RoleFactory(
            code="tpl-editor",
            permissions=_scoped(
                "references.templates.view",
                "references.templates.edit",
            ),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        resp = self._upload(api_client)
        assert resp.status_code == status.HTTP_201_CREATED

    def test_missing_file_returns_400(self, api_client):
        _superuser(api_client)
        resp = api_client.post(self.url, {}, format="multipart")
        assert resp.status_code == status.HTTP_400_BAD_REQUEST

    def test_unsupported_mime_returns_415(self, api_client):
        """Text files, PDFs, executables — anything outside the small
        image whitelist is rejected before touching storage."""
        _superuser(api_client)
        resp = self._upload(
            api_client,
            content=b"not an image",
            content_type="text/plain",
            name="naughty.txt",
        )
        assert resp.status_code == status.HTTP_415_UNSUPPORTED_MEDIA_TYPE
        assert resp.data["code"] == "unsupported_media_type"

    def test_oversized_file_returns_400(self, api_client):
        """5 MB limit. A single byte over the threshold is enough to
        trigger the gate — no partial acceptance."""
        _superuser(api_client)
        big = b"\x00" * (5 * 1024 * 1024 + 1)
        resp = self._upload(
            api_client,
            content=big,
            content_type="image/png",
            name="huge.png",
        )
        assert resp.status_code == status.HTTP_400_BAD_REQUEST
        assert resp.data["code"] == "file_too_large"

    def test_svg_accepted(self, api_client):
        """SVG is allowed — WeasyPrint renders it natively and the
        editor shows it as a static image (no script execution)."""
        _superuser(api_client)
        svg = b'<svg xmlns="http://www.w3.org/2000/svg" width="1" height="1"/>'
        resp = self._upload(
            api_client,
            content=svg,
            content_type="image/svg+xml",
            name="icon.svg",
        )
        assert resp.status_code == status.HTTP_201_CREATED
        assert resp.data["url"].endswith(".svg")


# --- DOCX validator + render -------------------------------------------


def _make_docx_bytes(paragraphs: list[str]) -> bytes:
    """Build a minimal valid .docx containing the given paragraphs.

    Used by validator and render tests — real .docx would be authored in
    Word, but for tests we synthesise via python-docx so the fixtures
    stay in-repo without committing binary blobs.
    """
    import io
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestDocxValidator:
    """`apps.contracts.services.docx_validator.validate` — extracts
    Jinja2 tags from a .docx and splits by known/unknown roots."""

    def test_extracts_dotted_paths(self):
        from apps.contracts.services import docx_validator
        data = _make_docx_bytes([
            "Договор №{{ contract.contract_number }} от {{ contract.date }}.",
            "Клиент: {{ client.full_name }}",
        ])
        res = docx_validator.validate(io.BytesIO(data))
        assert "contract.contract_number" in res.known
        assert "contract.date" in res.known
        assert "client.full_name" in res.known
        assert res.unknown == []
        assert res.as_dict()["is_valid"] is True

    def test_flags_unknown_root(self):
        from apps.contracts.services import docx_validator
        data = _make_docx_bytes([
            "{{ clinet.full_name }}",  # typo → unknown root
            "{{ contract.date }}",
        ])
        res = docx_validator.validate(io.BytesIO(data))
        assert res.known == ["contract.date"]
        assert res.unknown == ["clinet.full_name"]
        assert res.as_dict()["is_valid"] is False

    def test_for_loop_vars_are_not_flagged(self):
        """`{% for row in schedule %} {{ row.amount }} {% endfor %}` —
        `row` is a loop-local, must not land in unknown even though
        `row` isn't a known root."""
        from apps.contracts.services import docx_validator
        data = _make_docx_bytes([
            "{% for row in contract.schedule %}{{ row.amount }}{% endfor %}",
        ])
        res = docx_validator.validate(io.BytesIO(data))
        assert "contract.schedule" in res.known
        # `row.amount` is loop-local; shouldn't surface at all.
        assert all("row" not in t.split(".")[0] for t in res.all_tags)

    def test_qr_tag_is_known(self):
        """Pre-injected InlineImage root `qr` must validate."""
        from apps.contracts.services import docx_validator
        data = _make_docx_bytes(["{{ qr }}"])
        res = docx_validator.validate(io.BytesIO(data))
        assert res.known == ["qr"]

    def test_filter_pipe_tolerated(self):
        from apps.contracts.services import docx_validator
        data = _make_docx_bytes(["{{ client.full_name | upper }}"])
        res = docx_validator.validate(io.BytesIO(data))
        assert res.known == ["client.full_name"]


@pytest.mark.django_db
class TestValidateDocxEndpoint:
    """`POST /contract-templates/validate-docx/` — multipart upload,
    auth-gated like the image upload endpoint."""

    url = reverse("contract-template-validate-docx")

    def _post(self, api_client, content=None, content_type=None, name="tpl.docx"):
        from django.core.files.uploadedfile import SimpleUploadedFile
        if content is None:
            content = _make_docx_bytes(["{{ contract.contract_number }}"])
        ct = content_type or (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        f = SimpleUploadedFile(name, content, content_type=ct)
        return api_client.post(self.url, {"file": f}, format="multipart")

    def test_happy_path_returns_split(self, api_client):
        _superuser(api_client)
        resp = self._post(api_client)
        assert resp.status_code == status.HTTP_200_OK, resp.data
        assert resp.data["is_valid"] is True
        assert "contract.contract_number" in resp.data["known"]

    def test_unauthorized_blocked(self, api_client):
        role = RoleFactory(
            code="view-only-docx",
            permissions=_scoped("references.templates.view"),
        )
        api_client.force_authenticate(StaffFactory(role=role, password="x12345678"))
        resp = self._post(api_client)
        assert resp.status_code == status.HTTP_403_FORBIDDEN

    def test_wrong_mime_rejected(self, api_client):
        _superuser(api_client)
        resp = self._post(
            api_client,
            content=b"not a docx",
            content_type="text/plain",
            name="foo.txt",
        )
        assert resp.status_code == status.HTTP_415_UNSUPPORTED_MEDIA_TYPE

    def test_missing_file_400(self, api_client):
        _superuser(api_client)
        resp = api_client.post(self.url, {}, format="multipart")
        assert resp.status_code == status.HTTP_400_BAD_REQUEST


@pytest.mark.django_db
class TestDocxTemplateRender:
    """End-to-end: DOCX-sourced template → docxtpl fills tags → LibreOffice
    converts to PDF → bytes saved on Contract.file.

    The LibreOffice conversion step is mocked so the test doesn't spawn
    a subprocess — we only assert that docxtpl filled tags correctly
    and that the pipeline plumbed the right bytes to Contract.file.
    """

    def test_render_dispatches_to_docx_path(self, monkeypatch):
        from django.core.files.base import ContentFile
        from docx import Document
        from apps.contracts.services import docgen

        # Capture the filled .docx that docxtpl produced so we can
        # inspect its text content — bypassing LibreOffice entirely
        # keeps the test fast and deterministic.
        captured: dict[str, bytes] = {}

        def fake_convert(src_docx, *, profile_dir):
            with open(src_docx, "rb") as f:
                captured["docx"] = f.read()
            return b"%PDF-FAKE-BYTES"

        monkeypatch.setattr(docgen, "_docx_to_pdf", fake_convert)

        contract = ContractFactory(contract_number="DOCX-001")
        docx_bytes = _make_docx_bytes([
            "Договор №{{ contract.contract_number }} от {{ contract.date }}.",
        ])
        template = ContractTemplateFactory(
            source=ContractTemplate.Source.DOCX,
            file=ContentFile(docx_bytes, name="template.docx"),
            body="",
        )
        contract.template = template
        contract.save(update_fields=["template"])

        result = docgen.generate_pdf(contract)
        contract.refresh_from_db()

        assert result.pdf_size == len(b"%PDF-FAKE-BYTES")
        assert contract.file, "Contract.file must be populated"
        # Document snapshot records the source so audit knows which
        # pipeline produced this PDF.
        assert contract.document["source"] == "docx"

        # Unpack the captured filled .docx and inspect paragraph text
        # — proves docxtpl ran Jinja2 substitution end-to-end before
        # the (mocked) LibreOffice step.
        filled_doc = Document(io.BytesIO(captured["docx"]))
        text = "\n".join(p.text for p in filled_doc.paragraphs)
        assert "DOCX-001" in text
        assert "{{" not in text

    def test_missing_file_raises(self):
        from apps.contracts.services import docgen

        contract = ContractFactory()
        contract.template = ContractTemplateFactory(
            source=ContractTemplate.Source.DOCX,
            file=None,
            body="",
        )
        contract.save(update_fields=["template"])
        with pytest.raises(docgen.TemplateNotSet):
            docgen.generate_pdf(contract)
